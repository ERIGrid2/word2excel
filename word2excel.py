import math
import re
import sys
import os
import argparse
import shutil
from docx import Document
import docx
from docx.shape import InlineShape
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.reader import excel
from openpyxl.worksheet.worksheet import Worksheet

def is_test_case(table):
    cell = table.cell(0, 0)
    if cell.text.strip().lower() == 'name of the test case':
        return True
    return False

def is_test_specification(table):
    cell = table.cell(1, 0)
    if cell.text.strip().lower() == 'title of test':
        return True
    return False

def is_experiment_specification(table):
    cell = table.cell(1, 0)
    if cell.text.strip().lower() == 'title of experiment':
        return True
    return False

def get_inline_graphics(word_part, document):
    try:
        element = word_part._element
    except:
        return []

    drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

    graphics = []
    for drawing in element.findall('*//w:drawing', namespaces=element.nsmap):
        namespaces = element.nsmap
        if drawing_ns not in namespaces.values():
            namespaces['a'] = drawing_ns
        blip = drawing.find('*//a:blip[@r:embed]', namespaces=namespaces)
        if blip is not None:
            graphic = {}
            graphic_id = blip.embed
            image_part = document.part.related_parts[graphic_id]                
            graphic['name'] = os.path.basename(image_part.partname)
            graphic['data'] = image_part._blob
            graphics.append(graphic)
    
    for imagedata in element.findall('*//v:imagedata', namespaces=document._element.nsmap):
        graphic = {}
        image_id = imagedata.get('{{{0}}}id'.format(imagedata.nsmap['r']))
        image_part = document.part.related_parts[image_id]                
        graphic['name'] = os.path.basename(image_part.partname)
        graphic['data'] = image_part._blob
        graphics.append(graphic)
        
    return graphics

TEST_CASE_HEADLINE_REGEX = re.compile('Test Case\s(.*)')
TEST_SPECIFICATION_HEADLINE_REGEX = re.compile('Test Specification\s(.*)')
EXPERIMENT_SPECIFICATION_HEADLINE_REGEX = re.compile('Experiment Specification\s(.*)')

def is_bold(paragraph):
    for r in paragraph.runs:
        if not r.bold:
            return False
    return True 

def is_test_case_headline(paragraph):
    text = paragraph.text
    if TEST_CASE_HEADLINE_REGEX.match(text):
        if is_bold(paragraph):
            return True
    return False

def is_test_specification_headline(paragraph):
    text = paragraph.text
    if TEST_SPECIFICATION_HEADLINE_REGEX.match(text):
        if is_bold(paragraph):
            return True
    return False

def is_experiment_specification_headline(paragraph):
    text = paragraph.text
    if EXPERIMENT_SPECIFICATION_HEADLINE_REGEX.match(text):
        if is_bold(paragraph):
            return True
    return False

def is_qualification_strategy_headline(paragraph):
    text = paragraph.text
    if text.strip() == 'Qualification Strategy':
        if is_bold(paragraph):
            return True
    return False

def is_mapping_headline(paragraph):
    text = paragraph.text
    if text.strip() == 'Mapping to Research Infrastructure':
        if is_bold(paragraph):
            return True
    return False

def parse_test_case(table, document):
    test_case = {}

    re_author_version = re.compile('Author:?\s+(.*)\s+Version:?\s+(.*)')
    re_project_date = re.compile('Project:?\s+(.*)\s+Date:?\s+(.*)')
    re_qs = re.compile('Qualification Strategy')
    is_qs = False
    for p in document.paragraphs:
        text = p.text
        if is_test_case_headline(p):
            test_case['ID'] = {'desc': TEST_CASE_HEADLINE_REGEX.match(text).group(1).strip()} 
        if re_author_version.match(text):
            test_case['Author'] = {'desc': re_author_version.match(text).group(1).strip()}
            test_case['Version'] = {'desc': re_author_version.match(text).group(2).strip()}
        if re_project_date.match(text):
            test_case['Project'] = {'desc': re_project_date.match(text).group(1).strip()}
            test_case['Date'] = {'desc': re_project_date.match(text).group(2).strip()}
        
        if is_test_specification_headline(p):
            break
        if is_qualification_strategy_headline(p):
            is_qs = True
            test_case['Qualification Strategy'] = {'desc': '', 'graphics': []}
        elif is_qs:
            tc = test_case['Qualification Strategy']
            tc['desc'] = tc['desc']  + '\n' + get_paragraph_text(p) if tc['desc'] else get_paragraph_text(p)
            graphics = get_inline_graphics(p, document)
            if len(graphics) > 0:
                tc['graphics'].extend(graphics)

    for r, row in enumerate(table.rows):
        if len(row.cells) == 3:
            id = row.cells[1].text.split('\n')[0]
            id = id.split(':')[0]
            test_case[id.strip()] = {'desc': get_text(row.cells[2])}
            graphics = get_inline_graphics(row.cells[2], document)
            if len(graphics) > 0:
                test_case[id.strip()]['graphics'] = graphics

            # for p in row.cells[2].paragraphs:
            #     if is_bullet_list(p):
            #         print('- ' + p.text + ' (' + p.style.name + ')')
            #     if is_numbered_list(p):
            #         print('1. ' + p.text + ' (' + p.style.name + ')')

    return test_case

def get_text(cell):
    text = '\n'.join(get_paragraph_text(p) for p in cell.paragraphs)
    
    # See if there is a table within the cell that has text
    table_texts = '\n'.join(get_table_text(t) for t in cell.tables)

    return '\n'.join(t for t in [text, table_texts] if t)

def get_table_text(table):
    text = '\n'.join(get_text(c) for c in table._cells)
    return text

def get_paragraph_text(paragraph):
    prefix = ''
    if is_bullet_list(paragraph):
        level = get_numbering_level(paragraph)
        prefix = '    '.join(['' for _ in range(level)]) + '- '
    elif is_numbered_list(paragraph):
        level = get_numbering_level(paragraph)
        prefix = '    '.join(['' for _ in range(level)]) + '1. '
    return prefix + paragraph.text

def is_bullet_list(paragraph):
    return get_numbering_format(paragraph) == 'bullet'

def is_numbered_list(paragraph):
    fmt = get_numbering_format(paragraph)
    return fmt is not None and fmt != 'bullet'

def get_numbering_level(paragraph):
    lvl = get_numbering_lvl(paragraph)
    if lvl is not None:
        try:
            lvl_indent = lvl.find('w:pPr/w:ind', namespaces=lvl.nsmap)
            if lvl_indent is not None:
                level = math.floor(lvl_indent.left / 500000) + 1
            else:
                level = int(get_value_of_attribute(lvl, 'ilvl')) + 1

            return level
        except:
            pass
    return 0

def get_numbering_lvl(paragraph):
    document_part = paragraph.part
    namespaces = paragraph._element.nsmap
    w_namespace = namespaces['w']
    p_numbering = paragraph._element.find('*/w:numPr', namespaces=namespaces)
    if p_numbering is not None:
        att_val = '{' + w_namespace + '}val'
        ilvl = p_numbering.find('w:ilvl', namespaces=namespaces)
        numId = p_numbering.find('w:numId', namespaces=namespaces)
        if ilvl is not None and numId is not None:
            abstract_num_id = document_part.numbering_part.element.find(
                'w:num[@w:numId="' + get_attr_val(numId) + '"]/w:abstractNumId', 
                namespaces=namespaces)
            if abstract_num_id is not None:
                xpath_str = ('w:abstractNum[@w:abstractNumId="' + get_attr_val(abstract_num_id) + '"]' + 
                    '/w:lvl[@w:ilvl="' + get_attr_val(ilvl) + '"]')
                num_level = document_part.numbering_part.element.find(xpath_str, namespaces=namespaces)
                return num_level
    return None

def get_attr_val(element):
    return get_value_of_attribute(element, 'val')

def get_value_of_attribute(element, attribute_name):
    namespaces = element.nsmap
    w_namespace = namespaces['w']
    attribute = '{' + w_namespace + '}' + attribute_name
    return element.get(attribute)

def get_numbering_format(paragraph):
    num_level = get_numbering_lvl(paragraph)
    if num_level is not None:
        xpath_str = 'w:numFmt'
        num_fmt = num_level.find(xpath_str, namespaces=num_level.nsmap)
        if num_fmt is not None:
            return get_attr_val(num_fmt)
    return None

def find_test_specifications(document):
    test_specs = []
    
    is_mapping = False
    for p in document.paragraphs:
        text = p.text
        if is_test_specification_headline(p):
            test_spec = {}
            test_spec['ID'] = {'desc': TEST_SPECIFICATION_HEADLINE_REGEX.match(text).group(1).strip()}
            test_specs.append(test_spec)
        
        if is_experiment_specification_headline(p):
            break
        if is_mapping_headline(p):
            is_mapping = True
            test_spec = test_specs[-1]
            test_spec['Mapping to Research Infrastructure'] = {'desc': '', 'graphics': []}
        elif is_mapping:
            test_spec_mapping = test_specs[-1]['Mapping to Research Infrastructure']
            test_spec_mapping['desc'] = test_spec_mapping['desc'] + '\n' + get_paragraph_text(p) if test_spec_mapping['desc'] else get_paragraph_text(p)
            graphics = get_inline_graphics(p, document)
            if len(graphics) > 0:
                test_spec_mapping['graphics'].extend(graphics)
    return test_specs

def find_experiment_specifications(document):
    experiment_specs = []
    
    for p in document.paragraphs:
        text = p.text
        if is_experiment_specification_headline(p):
            exp_spec = {}
            exp_spec['ID'] = {'desc': EXPERIMENT_SPECIFICATION_HEADLINE_REGEX.match(text).group(1).strip()}
            experiment_specs.append(exp_spec)
    return experiment_specs

def parse_test_specification(table, document, test_spec):
    for r, row in enumerate(table.rows):
        if len(row.cells) == 2:
            id = row.cells[0].text.split('\n')[0]
            id = id.split(':')[0]
            test_spec[id.strip()] = {'desc': get_text(row.cells[1])}
            graphics = get_inline_graphics(row.cells[1], document)
            if len(graphics) > 0:
                test_spec[id.strip()]['graphics'] = graphics
    return test_spec    

def parse_experiment_specification(table, document, experiment_spec):
    for r, row in enumerate(table.rows):
        if len(row.cells) == 2:
            id = row.cells[0].text.split('\n')[0]
            id = id.split(':')[0]
            experiment_spec[id.strip()] = {'desc': get_text(row.cells[1])}
            graphics = get_inline_graphics(row.cells[1], document)
            if len(graphics) > 0:
                experiment_spec[id.strip()]['graphics'] = graphics
    return experiment_spec  

def write_diagrams(sheet, graphics):
    for i in range(1, 1000):
        if sheet['A' + str(i)].value == 'Diagrams':
            start_row = i + 1

    for i, graphic in enumerate(graphics):
        col = 3 + i
        sheet.cell(row=start_row, column=col).value = graphic['name']
        sheet.cell(row=start_row + 1, column=col).value = graphic['name']
        sheet.cell(row=start_row + 2, column=col).value = 'image'
        sheet.cell(row=start_row + 3, column=col).value = graphic['name']

def write_test_case(wb, test_case):
    test_case_sheet_template = wb['Test Case']
    test_case_sheet = wb.copy_worksheet(test_case_sheet_template)
    if 'ID' in test_case:
        test_case_sheet.title = test_case['ID']['desc']
    else:
        test_case_sheet.title = 'TC1'
    test_case_graphics = []
    for row in range(2, 1000):
        id_cell = test_case_sheet['B' + str(row)]
        if id_cell.value in test_case:
            value_cell = test_case_sheet['C' + str(row)]
            if type(value_cell.fill.fgColor.theme) == int:
                value_cell = test_case_sheet['C' + str(row + 1)]
            value_cell.value = test_case[id_cell.value]['desc']

            if 'graphics' in test_case[id_cell.value]:
                test_case_graphics.extend(test_case[id_cell.value]['graphics'])
                if test_case_sheet['B' + str(row + 2)].value == 'Diagram reference':
                    graphics_ref = '; '.join([g['name'] for g in test_case[id_cell.value]['graphics']])
                    test_case_sheet['C' + str(row + 2)].value = graphics_ref
    
    write_diagrams(test_case_sheet, test_case_graphics)

def write_test_specification(wb, test_spec):
    test_spec_sheet_template = wb['Test Specification']
    test_spec_sheet = wb.copy_worksheet(test_spec_sheet_template)
    if 'ID' in test_spec:
        test_spec_sheet.title = test_spec['ID']['desc']
    else:
        test_spec_sheet.title = 'TS1'
    test_spec_graphics = []
    for row in range(2, 1000):
        id_cell = test_spec_sheet['B' + str(row)]
        if id_cell.value in test_spec:
            value_cell = test_spec_sheet['C' + str(row)]
            if type(value_cell.fill.fgColor.theme) == int:
                value_cell = test_spec_sheet['C' + str(row + 1)]
            value_cell.value = test_spec[id_cell.value]['desc']

            if 'graphics' in test_spec[id_cell.value]:
                test_spec_graphics.extend(test_spec[id_cell.value]['graphics'])
                b_col_value = test_spec_sheet['B' + str(row + 2)].value
                if type(b_col_value) == str and b_col_value.lower() == 'diagram reference':
                    graphics_ref = '; '.join([g['name'] for g in test_spec[id_cell.value]['graphics']])
                    test_spec_sheet['C' + str(row + 2)].value = graphics_ref

    write_diagrams(test_spec_sheet, test_spec_graphics)

def write_experiment_specification(wb, test_spec):
    exp_spec_sheet_template = wb['Experiment Specification']
    exp_spec_sheet = wb.copy_worksheet(exp_spec_sheet_template)
    if 'ID' in test_spec:
        exp_spec_sheet.title = test_spec['ID']['desc']
    else:
        exp_spec_sheet.title = 'ES1'
    exp_spec_graphics = []
    for row in range(2, 1000):
        id_cell = exp_spec_sheet['B' + str(row)]
        if id_cell.value in test_spec:
            value_cell = exp_spec_sheet['C' + str(row)]
            if type(value_cell.fill.fgColor.theme) == int:
                value_cell = exp_spec_sheet['C' + str(row + 1)]
            value_cell.value = test_spec[id_cell.value]['desc']

            if 'graphics' in test_spec[id_cell.value]:
                exp_spec_graphics.extend(test_spec[id_cell.value]['graphics'])
                b_col_value = exp_spec_sheet['B' + str(row + 2)].value
                if type(b_col_value) == str and b_col_value.lower() == 'diagram reference':
                    graphics_ref = '; '.join([g['name'] for g in test_spec[id_cell.value]['graphics']])
                    exp_spec_sheet['C' + str(row + 2)].value = graphics_ref

    write_diagrams(exp_spec_sheet, exp_spec_graphics)

def word2excel(doc_filename, template_path, create_folder=False, copy_word_file=False):
    filepath = os.path.dirname(doc_filename)
    name_of_doc_file = '.'.join(os.path.basename(doc_filename).split('.')[:-1])

    document = None
    try:
        document = Document(doc_filename)
    except:
        print('ERROR: Could not open Word file: {0}'.format(doc_filename))
        return

    if create_folder:
        try:
            new_folder = os.path.join(filepath, name_of_doc_file)
            os.mkdir(new_folder)
        except FileExistsError as e:
            pass
        except:
            print('ERROR: Could not create folder: {0}'.format(new_folder))
            return
        filepath = new_folder

        if copy_word_file:
            dest_path = os.path.join(filepath, os.path.basename(doc_filename))
            try:
                shutil.copyfile(doc_filename, dest_path)
            except OSError:
                print("ERROR: Could not write to destination {0}.".format(dest_path))
                return

    test_case = {}
    test_specifications = find_test_specifications(document)
    experiment_specifications = find_experiment_specifications(document)

    # parse docx file
    number_test_specs = 0
    number_experiment_specs = 0
    for t, table in enumerate(document.tables):
        if is_test_case(table):
            test_case = parse_test_case(table, document)
        elif is_test_specification(table):
            if number_test_specs < len(test_specifications):
                parse_test_specification(table, document, test_specifications[number_test_specs])
            else:
                test_specifications.append(parse_test_specification(table, document, {}))
            number_test_specs += 1
        elif is_experiment_specification(table):
            if number_experiment_specs < len(experiment_specifications):
                parse_experiment_specification(table, document, experiment_specifications[number_experiment_specs])
            else:
                experiment_specifications.append(parse_experiment_specification(table, document, {}))
            number_experiment_specs += 1
            
    # write to excel file and save images
    excelfile = os.path.join(filepath, name_of_doc_file + '.xlsx')
    try:
        shutil.copyfile(template_path, excelfile)
    except FileNotFoundError:
        print("ERROR: Excel template {0} does not exist.".format(template_path))
        return
    except OSError:
        print("ERROR: Could not write to destination {0}.".format(excelfile))
        return

    try:
        wb = load_workbook(excelfile)
    except:
        print('ERROR: Could not open Excel file: {0}'.format(excelfile))
        return

    write_test_case(wb, test_case)
    for test_spec in test_specifications:
        write_test_specification(wb, test_spec)
    for exp_spec in experiment_specifications:
        write_experiment_specification(wb, exp_spec)
    wb.save(excelfile)
    
    # save images
    for part in document.part.related_parts.values():
        if type(part) == docx.ImagePart:
            image_part = part
            
            image_name = os.path.basename(image_part.partname)
            image_path = os.path.join(filepath, image_name)
        
            with open(image_path, 'wb') as fs:
                fs.write(image_part._blob)

if __name__ == '__main__':
    excel_template_default = './template/HTD_TEMPLATE_V1.2.xlsx'

    parser = argparse.ArgumentParser(description='Converts test cases according to the ERIGrid HTD Template from Word into Excel files.')
    parser.add_argument('path', help='Path to either a Word file or a folder. If a folder is provided, all Word files in that folder will be converted.')
    parser.add_argument('-t', '--excel-template', help='Path to the Excel template that should be used. Standard: {0}'.format(excel_template_default),
                        default=os.path.join(os.path.dirname(os.path.abspath(__file__)), excel_template_default))
    parser.add_argument('-f', '--create-folder', help='Saves the Excel file and extracted images to a folder with the name of Word file.', 
                        action='store_true')
    parser.add_argument('-c', '--copy-word-file', help='Copies the Word file into the new folder', action='store_true')
    args = parser.parse_args()    

    doc_filename = args.path
    template_path = args.excel_template
    create_folder = args.create_folder
    copy_word_file = args.copy_word_file

    files_to_convert = []

    if os.path.isdir(doc_filename):
        for f in os.scandir(doc_filename):
            if f.is_file and f.path.endswith('.docx'):
                files_to_convert.append(f.path)
    else:
        files_to_convert.append(doc_filename)
        
    for f in files_to_convert:
        print('\nConverting {0}'.format(f))
        word2excel(f, template_path, create_folder=create_folder, copy_word_file=copy_word_file)